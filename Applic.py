# Импорт библиотек
# GUI
from PyQt5.QtCore import *
from PyQt5.QtGui import *
from PyQt5.QtWidgets import *
import sys
import math
import sqlite3
from itertools import groupby

class Slice(QMainWindow):
    def __init__(self):
        super().__init__()
        self.setWindowTitle("Расчёт на срез")
        self.setFixedSize(780,420)

        self.con = sqlite3.connect('MaterialT.db')
        self.cursor = self.con.cursor()
        self.cursor.execute('select MaterialGOST from MaterialGOST')
        self.materialsGost = self.cursor.fetchall()
        self.cursor.execute('select DiametrRezbi from RezbaMetricheskayaGOST')
        self.diamRezbiMetr = self.cursor.fetchall()
        self.cursor.execute('select Shag from RezbaMetricheskayaGOST')
        self.shagRezbiMetr = self.cursor.fetchall()
        self.cursor.execute("select Naprejenie from MaterialGOST WHERE MaterialGOST ='Ст2 ГОСТ 380-94'")
        self.materialsZnach = self.cursor.fetchall()
        self.pi = math.pi
        self.initUI()
        self.show()

    def initUI(self):
        self.tabWidget = QTabWidget(self)
        self.tabWidget.setGeometry(QRect(0, 0, 801, 601))
        self.tab1 = QWidget()
        self.tab2 = QWidget()
        # TAB 1
        self.tabWidget.addTab(self.tab1, "Тип крепления")
        self.tabWidget.addTab(self.tab2, "Вывод")
        self.typeMountC = QComboBox(self.tab1)
        self.typeMountC.setGeometry(QRect(20, 10, 291, 25))
        self.typeMountC.addItems(["Резьбовое метрическое","Резьбовое дюймовое","Заклепочное","Заклепочное, с полой заклепкой","Заклепочное, резьбовой заклепкой","Штифтовое продольное","Штифтовое поперечное","Призматической шпонкой","Сегментной шпонкой","Клиновой шпонкой","Шлицевое" ])
        self.mountingSize = QLabel('Диаметр резьбы',self.tab1)
        self.mountingSize.setGeometry(QRect(390, 10, 250, 25))
        self.mountingSizeC = QComboBox(self.tab1)
        self.sortDiamRezbiMetr = [el for el, _ in groupby(self.diamRezbiMetr)]
        for diamMetr  in self.sortDiamRezbiMetr:
            # diamMetr = str(diamMetr)
            self.mountingSizeC.addItems(diamMetr)
        self.mountingSizeC.setGeometry(QRect(500, 10, 100, 25))
        self.shagRezbi = QLabel('Шаг',self.tab1)
        self.shagRezbi.setGeometry(QRect(660, 10, 100, 25))
        self.shagRezbiC = QComboBox(self.tab1)
        self.shagRezbiC.addItem("0,075")
        self.shagRezbiC.setGeometry(QRect(690, 10, 60, 25))
        self.numberMounts = QLabel('Количество креплений ',self.tab1)
        self.numberMounts.setGeometry(QRect(20, 50, 200, 25))
        self.numberMountsE = QLineEdit('1',self.tab1)
        self.numberMountsE.setGeometry(QRect(210, 50, 100, 25))
        self.material = QLabel('Материал крепления',self.tab1)
        self.material.setGeometry(QRect(390, 50, 250, 25))
        self.materialC = QComboBox(self.tab1)
        for self.materialGost in self.materialsGost:
            self.materialC.addItems(self.materialGost)
        self.materialC.setGeometry(QRect(550, 50, 200, 25))
        # groupBox1
        self.groupBox1 = QGroupBox("Для резьбового и заклепочного соединений ", self.tab1)
        self.groupBox1.setGeometry(QRect(10, 110, 350, 264))
        self.diameter = QLabel('Диаметр отверстия',self.groupBox1)
        self.diameter.setGeometry(QRect(20, 25, 200, 25))
        self.diameterE = QLineEdit('23',self.groupBox1)
        self.diameterE.setGeometry(QRect(175, 25, 140, 25))
        self.plate = QLabel('Пластины',self.groupBox1)
        self.plate.setGeometry(QRect(20, 60, 150, 25))
        self.plateT= QTableWidget(self.groupBox1)
        self.plateT.setColumnCount(2)
        self.plateT.setRowCount(1)
        self.plateT.setHorizontalHeaderLabels(['Толщина пластины, мм  ','Поперечная сила, Н '])
        self.plateT.setGeometry(QRect(20, 80, 295, 170))
        self.plateT.resizeColumnsToContents()
        # groupBox2
        self.groupBox2 = QGroupBox("Вал", self.tab1)
        self.groupBox2.setGeometry(QRect(390, 110, 360, 220))
        self.shaftDiameter = QLabel('Диаметр вала, мм',self.groupBox2)
        self.shaftDiameter.setGeometry(QRect(20, 25, 200, 25))
        self.shaftDiameterE = QLineEdit('70',self.groupBox2)
        self.shaftDiameterE.setGeometry(QRect(175, 25, 140, 25))
        self.torque = QLabel('Крутящий момент в H*мм',self.groupBox2)
        self.torque.setGeometry(QRect(20, 65, 200, 25))
        self.torqueC = QLineEdit('166440',self.groupBox2)
        self.torqueC.setGeometry(QRect(175, 65, 140, 25))
        self.vistupShponki = QLabel('Выступ шпонки,мм',self.groupBox2)
        self.vistupShponki.setGeometry(QRect(20, 105, 200, 25))
        self.vistupShponkiE = QLineEdit('3',self.groupBox2)
        self.vistupShponkiE.setGeometry(QRect(175, 105, 140, 25))
        self.tolzinaShponki = QLabel('Толщина шпонки,мм', self.groupBox2)
        self.tolzinaShponki.setGeometry(QRect(20,145,200,25))
        self.tolzinaShponkiE = QLineEdit('10',self.groupBox2)
        self.tolzinaShponkiE.setGeometry(QRect(175,145,140,25))
        self.dlinaShponki = QLabel('Длина шпонки, мм', self.groupBox2)
        self.dlinaShponki.setGeometry(QRect(20,185,200,25))
        self.dlinaShponkiE = QLineEdit('40',self.groupBox2)
        self.dlinaShponkiE.setGeometry(QRect(175,185,140,25))

        self.calculation = QPushButton('Расчитать',self.tab1)
        self.calculation.setGeometry(QRect (391, 338, 100, 35))
        self.outputInW = QPushButton('Вывести',self.tab1)
        self.outputInW.setGeometry(QRect (505, 338, 100, 35))
        self.exitApp = QPushButton('Выход',self.tab1)
        self.exitApp.setGeometry(QRect (650, 338, 100, 35))

        self.shaftDiameter.setEnabled(False)
        self.shaftDiameterE.setEnabled(False)
        self.torque.setEnabled(False)
        self.torqueC.setEnabled(False)
        self.vistupShponki.setEnabled(False)
        self.vistupShponkiE.setEnabled(False)
        self.tolzinaShponki.setEnabled(False)
        self.tolzinaShponkiE.setEnabled(False)
        self.dlinaShponki.setEnabled(False)
        self.dlinaShponkiE.setEnabled(False)

        self.exitApp.pressed.connect(self.exitA)
        self.plateT.activated.connect(self.plusRow)
        self.calculation.pressed.connect(self.raschet)
        self.typeMountC.currentTextChanged.connect(self.changeForm)
        self.materialC.currentTextChanged.connect(self.materialChanged)
        self.mountingSizeC.currentTextChanged.connect(self.diametrChanged)
        self.typeMountC.currentTextChanged.connect(self.chengeDiam)
        self.outputInW.pressed.connect(self.outWord)
# FIX: ДОБАВИТЬ КНОПКУ ВЫВОДА:

        # 2 TAB :3
        self.tipRezania = QLabel('Тип резанья:', self.tab2)
        self.tipRezania.setGeometry(QRect(20, 10, 291, 25))
        self.tipRezaniaV = QLabel('', self.tab2)
        self.tipRezaniaV.setGeometry(QRect(170, 10, 291, 25))
        self.kolvoKrep = QLabel('Кол-во креплений:', self.tab2)
        self.kolvoKrep.setGeometry(QRect(20, 50, 291, 25))
        self.kolvoKrepV = QLabel('', self.tab2)
        self.kolvoKrepV.setGeometry(QRect(170, 50, 291, 25))
        self.diameterRezbi = QLabel('Диаметр резьбы:',self.tab2)
        self.diameterRezbi.setGeometry(QRect(450, 10, 200, 25))
        self.diameterRezbiV = QLabel('',self.tab2)
        self.diameterRezbiV.setGeometry(QRect(600, 10, 200, 25))
        self.diameterShag = QLabel('Шаг:',self.tab2)
        self.diameterShag.setGeometry(QRect(680, 10, 200, 25))
        self.diameterShagV = QLabel('',self.tab2)
        self.diameterShagV.setGeometry(QRect(730, 10, 200, 25))
        self.materialT2 = QLabel('Материал крепления:',self.tab2)
        self.materialT2.setGeometry(QRect(450, 50, 200, 25))
        self.materialT2V = QLabel('',self.tab2)
        self.materialT2V.setGeometry(QRect(620, 50, 200, 25))
        self.chooseYDest = QLabel('',self.tab2)
        self.chooseYDest.setGeometry(QRect(0, 90, 2000, 25))
        # Штуки введённые
        self.tableChoose1 = QLabel('',self.tab2)
        self.tableChoose1.setGeometry(QRect(20, 120, 200, 25))
        self.tableChoose1V = QLabel('',self.tab2)
        self.tableChoose1V.setGeometry(QRect(300, 120, 200, 25))
        self.tableChoose2 = QLabel('',self.tab2)
        self.tableChoose2.setGeometry(QRect(20, 160, 200, 25))
        self.tableChoose2V = QLabel('',self.tab2)
        self.tableChoose2V.setGeometry(QRect(300, 160, 200, 25))
        self.tableChoose3 = QLabel('',self.tab2)
        self.tableChoose3.setGeometry(QRect(20, 200, 200, 25))
        self.tableChoose3V = QLabel('',self.tab2)
        self.tableChoose3V.setGeometry(QRect(300, 200, 200, 25))
        self.tableChoose4 = QLabel('',self.tab2)
        self.tableChoose4.setGeometry(QRect(20, 240, 200, 25))
        self.tableChoose4V = QLabel('',self.tab2)
        self.tableChoose4V.setGeometry(QRect(300, 240, 200, 25))
        self.tableChoose5 = QLabel('',self.tab2)
        self.tableChoose5.setGeometry(QRect(20, 280, 200, 25))
        self.tableChoose5V = QLabel('',self.tab2)
        self.tableChoose5V.setGeometry(QRect(300, 280, 200, 25))
        # Формулы
        self.tableChooseRight1 = QLabel('Tcp =',self.tab2)
        self.tableChooseRight1.setGeometry(QRect(450, 120, 200, 25))
        self.tableChooseRight2 = QLabel('',self.tab2)
        self.tableChooseRight2.setGeometry(QRect(600,120,200,25))


    # Добавление сторик для таблицы
    @pyqtSlot()
    def plusRow(self):
        self.rowPosition = self.plateT.rowCount()
        self.plateT.insertRow(self.rowPosition)
    # Изменинее "Диаметр резьбы" и "шаг"
    @pyqtSlot()
    def diametrChanged(self):
        if self.typeMountC.currentText() == 'Резьбовое метрическое' or self.typeMountC.currentText() == 'Штифтовое продольное' or self.typeMountC.currentText() =='Штифтовое поперечное' or self.typeMountC.currentText() =='Призматической шпонкой' or self.typeMountC.currentText() =='Сегментной шпонкой' or self.typeMountC.currentText() =='Клиновой шпонкой' or self.typeMountC.currentText() == 'Шлицевое':
            self.shagRezbiC.clear()
            self.cursor.execute('select Shag from RezbaMetricheskayaGOST WHERE DiametrRezbi = ?' ,(self.mountingSizeC.currentText(),))
            self.metrShagi = self.cursor.fetchall()
            for self.metrShag  in self.metrShagi:
                self.shagRezbiC.addItems(self.metrShag)
        elif self.typeMountC.currentText() == 'Резьбовое дюймовое':
            self.shagRezbiC.clear()
            self.cursor.execute('select Shag from RezbaDuimovayaBoltGOST WHERE DiametrRezbi = ?' ,(self.mountingSizeC.currentText(),))
            self.duimShagi = self.cursor.fetchall()
            for self.duimShag  in self.duimShagi:
                self.shagRezbiC.addItems(self.duimShag)
        else:
            self.shagRezbiC.clear()
            self.cursor.execute('select Shag from DuimovayaRezbaZaklepok WHERE DiametrRezbi = ?' ,(self.mountingSizeC.currentText(),))
            self.duimZaklShagi = self.cursor.fetchall()
            for self.duimZaklShag  in self.duimZaklShagi:
                self.shagRezbiC.addItems(self.duimZaklShag)
    # Подсос материала
    @pyqtSlot()
    def materialChanged(self):
        self.gostMaterialCombo = self.materialC.currentText()
        self.cursor.execute('select Naprejenie from MaterialGOST WHERE MaterialGOST = ?' ,(self.gostMaterialCombo,))
        self.materialsZnach = self.cursor.fetchall()
    # Изменение формы заглушить вал и тд
    @pyqtSlot()
    def changeForm(self):
        if self.typeMountC.currentText() == "Резьбовое метрическое"or self.typeMountC.currentText() =="Резьбовое дюймовое"or self.typeMountC.currentText() =="Заклепочное"or self.typeMountC.currentText() =="Заклепочное, с полой заклепкой"or self.typeMountC.currentText() =="Заклепочное, резьбовой заклепкой":
            self.shaftDiameter.setEnabled(False)
            self.shaftDiameterE.setEnabled(False)
            self.torque.setEnabled(False)
            self.torqueC.setEnabled(False)
            self.vistupShponki.setEnabled(False)
            self.vistupShponkiE.setEnabled(False)
            self.tolzinaShponki.setEnabled(False)
            self.tolzinaShponkiE.setEnabled(False)
            self.dlinaShponki.setEnabled(False)
            self.dlinaShponkiE.setEnabled(False)
            self.diameter.setEnabled(True)
            self.diameterE.setEnabled(True)
            self.plate.setEnabled(True)
            self.plateT.setEnabled(True)
        else:
            self.shaftDiameter.setEnabled(True)
            self.shaftDiameterE.setEnabled(True)
            self.torque.setEnabled(True)
            self.torqueC.setEnabled(True)
            self.vistupShponki.setEnabled(True)
            self.vistupShponkiE.setEnabled(True)
            self.tolzinaShponki.setEnabled(True)
            self.tolzinaShponkiE.setEnabled(True)
            self.dlinaShponki.setEnabled(True)
            self.dlinaShponkiE.setEnabled(True)
            self.diameter.setEnabled(False)
            self.diameterE.setEnabled(False)
            self.plate.setEnabled(False)
            self.plateT.setEnabled(False)
    @pyqtSlot()
    def chengeDiam(self):
        if self.typeMountC.currentText() == 'Резьбовое метрическое' or self.typeMountC.currentText() == 'Штифтовое продольное' or self.typeMountC.currentText() =='Штифтовое поперечное' or self.typeMountC.currentText() =='Призматической шпонкой' or self.typeMountC.currentText() =='Сегментной шпонкой' or self.typeMountC.currentText() =='Клиновой шпонкой' or self.typeMountC.currentText() == 'Шлицевое':
            self.mountingSizeC.clear()
            self.cursor.execute('select DiametrRezbi from RezbaMetricheskayaGOST')
            self.changeMDiametri=self.cursor.fetchall()
            self.sortMChangeDiametri = [el for el, _ in groupby(self.changeMDiametri)]
            for self.sortMChangeDiametr in self.sortMChangeDiametri:
                self.mountingSizeC.addItems(self.sortMChangeDiametr)
        elif self.typeMountC.currentText() == 'Резьбовое дюймовое':
            self.mountingSizeC.clear()
            self.cursor.execute('select DiametrRezbi from RezbaDuimovayaBoltGOST')
            self.changeDDiametri=self.cursor.fetchall()
            self.sortDChangeDiametri = [el for el, _ in groupby(self.changeDDiametri)]
            for self.sortDChangeDiametr in self.sortDChangeDiametri:
                self.mountingSizeC.addItems(self.sortDChangeDiametr)
        else:
            self.mountingSizeC.clear()
            self.cursor.execute('select DiametrRezbi from DuimovayaRezbaZaklepok')
            self.duimZaklepki=self.cursor.fetchall()
            self.sortduimZaklepki = [el for el, _ in groupby(self.duimZaklepki)]
            for self.sortduimZaklepk in self.sortduimZaklepki:
                self.mountingSizeC.addItems(self.sortduimZaklepk)
    @pyqtSlot()
    def raschet(self):
        self.tipRezaniaV.setText(self.typeMountC.currentText())
        self.kolvoKrepV.setText(self.numberMountsE.text())
        self.diameterRezbiV.setText(self.mountingSizeC.currentText())
        self.diameterShagV.setText(self.shagRezbiC.currentText())
        self.materialT2V.setText(self.materialC.currentText())
        if self.diameterE.isEnabled():
            # --------------------------------------------
            p = self.materialsZnach
            i = self.plateT.rowCount()
            z = self.numberMountsE.text()
            p = list(map(tuple, p))
            p = p[0]
            p = str(p)
            p = p.replace(',', '')
            p = p.replace(')','')
            p = p.replace('(','')
            p = int(p)
            self.Q = (p/(int(i)*int(z)))
            # ---------------------------------------------------
            try:
                d = self.mountingSizeC.currentText()
                d = d.replace(',', '.')
                d = float(d)
            except ValueError:
                d = 10
            self.Fcp = (round(int(self.pi))) *(d*d)/4
            TcpF = self.Q/ self.Fcp
            print(TcpF)

            self.chooseYDest.setText('----------------------------------------------------Для резьбового и заклепочного соединений-----------------------------------------------------------')
            self.tableChoose1.setText('Диаметр отверстия:')
            self.tableChoose1V.setText(self.diameterE.text())
            self.tableChoose2.setText('Количество пластин:')
            self.tableChoose2V.setText(str(self.plateT.rowCount()))
            self.tableChoose3.setText('Толщина пластин:')
            self.tableChooseRight2.setText(str(TcpF))
            massTolz = []
            try:
                for i in range(self.plateT.rowCount()):
                    massTolz.append(self.plateT.item(i, 0).text())
                    print(massTolz)
                    massTolzR = list(map(int, massTolz))
                    try:
                        massTolzRe = sum(massTolzR)
                        print(massTolzRe)
                        self.tableChoose3V.setText(str(massTolzRe))
                    except TypeError:
                        self.tableChoose3V.setText(self.plateT.item(0, 0).text())
            except AttributeError:
                self.tableChoose3V.setText("нет")

            self.tableChoose4.setText('Поперечная сила:')
            self.tableChoose5.setText('')
            self.tableChoose5V.setText('')
            poperchSila =[]
            try:
                for i in range(self.plateT.rowCount()):
                    poperchSila.append(self.plateT.item(i,1).text())
                    print(poperchSila)
                    poperchSilaR = list(map(int, poperchSila))
                    try:
                        poperchSilaRe = sum(poperchSilaR)
                        print(poperchSilaRe)
                        self.tableChoose4V.setText(str(poperchSilaRe))
                    except TypeError:
                        self.tableChoose4V.setText(self.plateT.item(0,1).text())
            except AttributeError:
                self.tableChoose4V.setText("нет")

        # Расчёт вала
        else:
            d = self.shaftDiameterE.text()
            k = self.vistupShponkiE.text()
            b = self.tolzinaShponkiE.text()
            l = self.dlinaShponkiE.text()
            M = self.torqueC.text()
            TpcS = int(M)/((0.5*(int(d)+int(k)))*int(b)*int(l))
            print(TpcS)
            self.chooseYDest.setText('-----------------------------------------------------------------------Для Вала------------------------------------------------------------------------------')
            self.tableChooseRight2.setText(str(TpcS))
            self.tableChoose1.setText('Диаметр вала')
            self.tableChoose2.setText('Крутящий момент')
            self.tableChoose3.setText('Выступ шпонки')
            self.tableChoose4.setText('Толщина шпонки')
            self.tableChoose5.setText('Длина шпонки')
            self.tableChoose1V.setText(self.shaftDiameterE.text())
            self.tableChoose2V.setText(self.torqueC.text())
            self.tableChoose3V.setText(self.vistupShponkiE.text())
            self.tableChoose4V.setText(self.tolzinaShponkiE.text())
            self.tableChoose5V.setText(self.dlinaShponkiE.text())
    @pyqtSlot()
    def outWord(self):
        import docx
        from docx import Document
        document = Document()
        document.add_heading('Расчёт на срез',0)
        par1 = document.add_paragraph('Тип резанья:  ')
        par1.add_run(self.typeMountC.currentText())
        par2 = document.add_paragraph('Кол-во креплений:  ')
        par2.add_run(self.numberMountsE.text())
        par3 = document.add_paragraph('Диаметр резьбы:  ')
        par3.add_run(self.mountingSizeC.currentText())
        par4 = document.add_paragraph('Шаг:  ')
        par4.add_run(self.shagRezbiC.currentText())
        par5 = document.add_paragraph('Материал крепления:  ')
        par5.add_run(self.materialC.currentText())
        if self.diameterE.isEnabled():
            # --------------------------------------------
            p = self.materialsZnach
            i = self.plateT.rowCount()
            z = self.numberMountsE.text()
            p = list(map(tuple, p))
            p = p[0]
            p = str(p)
            p = p.replace(',', '')
            p = p.replace(')','')
            p = p.replace('(','')
            p = int(p)
            self.Q = (p/(int(i)*int(z)))
            # ---------------------------------------------------
            try:
                d = self.mountingSizeC.currentText()
                d = d.replace(',', '.')
                d = float(d)
            except ValueError:
                d = 10
            self.Fcp = (round(int(self.pi))) *(d*d)/4
            TcpF = self.Q/ self.Fcp
            print(TcpF)
            document.add_paragraph('Расчёт проводился для резьбового и заклепочного соединений')
            par6 = document.add_paragraph('Диаметр отверстия:  ')
            par6.add_run(self.diameterE.text())
            par7 =  document.add_paragraph('Количество пластин:  ')
            par7.add_run(str(self.plateT.rowCount()))
            par8 =  document.add_paragraph('Толщина пластин:  ')
            massTolz = []
            try:
                for i in range(self.plateT.rowCount()):
                    massTolz.append(self.plateT.item(i, 0).text())
                    massTolzR = list(map(int, massTolz))
                    try:
                        massTolzRe = sum(massTolzR)
                        par8.add_run(str(massTolzRe))
                    except TypeError:
                        par8.add_run(self.plateT.item(0, 0).text())
            except AttributeError:
                par8.add_run('Нет')
            par10 = document.add_paragraph('Поперечная сила:  ')
            poperchSila =[]
            try:
                for i in range(self.plateT.rowCount()):
                    poperchSila.append(self.plateT.item(i,1).text())
                    print(poperchSila)
                    poperchSilaR = list(map(int, poperchSila))
                    try:
                        poperchSilaRe = sum(poperchSilaR)
                        print(poperchSilaRe)
                        par10.add_run(str(poperchSilaRe))
                    except TypeError:
                        par10.add_run(self.plateT.item(0,1).text())
            except AttributeError:
                par10.add_run("Нет")
            par9 =  document.add_paragraph('Tcp = ')
            par9.add_run(str(TcpF))
        else:
            d = self.shaftDiameterE.text()
            k = self.vistupShponkiE.text()
            b = self.tolzinaShponkiE.text()
            l = self.dlinaShponkiE.text()
            M = self.torqueC.text()
            TpcS = int(M)/((0.5*(int(d)+int(k)))*int(b)*int(l))
            print(TpcS)
            document.add_paragraph('Расчёт проводился для Вала')
            par6=document.add_paragraph('Диаметр вала:  ')
            par7=document.add_paragraph('Крутящий момент:  ')
            par8=document.add_paragraph('Выступ шпонки:  ')
            par9=document.add_paragraph('Толщина шпонки:  ')
            par10=document.add_paragraph('Длина шпонки:  ')
            par6.add_run(self.shaftDiameterE.text())
            par7.add_run(self.torqueC.text())
            par8.add_run(self.vistupShponkiE.text())
            par9.add_run(self.tolzinaShponkiE.text())
            par10.add_run(self.dlinaShponkiE.text())
            fin=document.add_paragraph('Tpc = ')
            fin.add_run(str(TpcS))
        document.save('Расчёт на срез.docx')
    # выход
    @pyqtSlot()
    def exitA(self):
        self.con.close()
        sys.exit()


if __name__ == '__main__':
    # для списоков аргументов командной строки
    App = QApplication(sys.argv)
    # Конструктор
    window = Slice()
    # Реакция на крестик
    sys.exit(App.exec())
